from docx import Document
import sys

document = Document()
res = ' '.join(i + '\n' for i in sys.stdin)
for i in res:
    document.add_heading('Հրավիրատոմս\n', 0)
    p = document.add_paragraph(f'Հարգելի {i}\n')
    p.bold = True
    p.add_run('Մենք սպասում ենք ձեզ մարտի 8-ի կապակցությամբ մեր հանդիսությունների դահլիճում\n').bold = True
    p.add_run('12:30\n')
    document.add_page_break()

'''document.add_heading('Առաջին մակարդակի վերնագիրը', level=1)
document.add_paragraph('Ինչ-որ մեջբերում', style='Intense Quote')

document.add_paragraph('Չհամարակալած ցուցակի տարր',
                       style='List Bullet')
document.add_paragraph('Համարակալած ցուցակի տարր',
                       style='List Number')

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Համար'
hdr_cells[1].text = 'Անվանում'
hdr_cells[2].text = 'Քանակ'''

document.save('test.docx')